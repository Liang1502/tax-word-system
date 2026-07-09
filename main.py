from fastapi import FastAPI, Query
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from pydantic import BaseModel
from docx import Document
import hashlib
import json
import re
import uuid
import time
import asyncio
import os

try:
    import psycopg
    from psycopg.rows import dict_row
except Exception:
    psycopg = None
    dict_row = None

app = FastAPI()

TEMPLATE_PATH = "templates/納保申請書_官方格式_可套填_v8.docx"

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)
FEEDBACK_PATH = os.path.join(OUTPUT_DIR, "feedback.jsonl")
DATABASE_URL = os.getenv("DATABASE_URL", "").strip()

download_store = {}

# 追蹤 /form URL 單次使用狀態：hash → expire_time（首次存取後 30 分鐘內封鎖）
form_access_store: dict[str, float] = {}

EXPIRED_HTML = """<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>連結已失效</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Noto Sans TC',sans-serif;background:#f0f4f8;display:flex;align-items:center;justify-content:center;min-height:100vh}
.box{background:#fff;border-radius:12px;padding:2rem 2.5rem;box-shadow:0 2px 12px rgba(0,0,0,.1);text-align:center;max-width:480px}
h1{color:#c0392b;font-size:1.15rem;margin-bottom:1rem}
p{color:#555;font-size:.93rem;line-height:1.7}
</style>
</head>
<body>
<div class="box">
  <h1>⚠️ 此連結已失效</h1>
  <p>每個申請書產製連結僅能使用一次。<br>如需重新產製，請返回對話重新操作。</p>
</div>
</body>
</html>"""


class GenerateRequest(BaseModel):

    apply_year: str
    apply_month: str
    apply_day: str

    formal_statement: str

    case_category_suggestion: str
    tax_items_suggestion: str

    apply_method_suggestion: str
    reply_method_suggestion: str
    notify_method_suggestion: str

    notify_email: str = ""

    representative_name: str = ""
    representative_address: str = ""
    representative_phone: str = ""

    agent_name: str = ""
    agent_address: str = ""
    agent_phone: str = ""

    evidence_list: str = ""


class FeedbackRequest(BaseModel):

    rating: int
    comment: str = ""
    stage: str = ""
    case_category: str = ""
    tax_item: str = ""


def replace_text(paragraph, replacements):

    text = paragraph.text

    for key, value in replacements.items():
        if key in text:
            text = text.replace(key, value)

    paragraph.text = text


def replace_all(doc, replacements):

    for p in doc.paragraphs:
        replace_text(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p, replacements)


def remove_email_label(doc):

    targets = ["{{notify_email}}", "信箱", "Email", "（信箱：）"]

    for p in doc.paragraphs:
        text = p.text
        for t in targets:
            text = text.replace(t, "")
        p.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for t in targets:
                        text = text.replace(t, "")
                    p.text = text


def has_placeholder(doc):

    pattern = re.compile(r"\{\{.*?\}\}")

    for p in doc.paragraphs:
        if pattern.search(p.text):
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if pattern.search(p.text):
                        return True

    return False


def normalize_choice(value: str, mapping: dict[str, str]) -> str:
    value = (value or "").strip()
    return mapping.get(value, value)


def write_feedback(data: FeedbackRequest):
    rating = int(data.rating)
    if rating < 1 or rating > 5:
        raise ValueError("rating must be between 1 and 5")

    record = {
        "id": str(uuid.uuid4()),
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S%z", time.localtime()),
        "rating": rating,
        "comment": (data.comment or "").strip()[:1000],
        "stage": (data.stage or "").strip()[:40],
        "case_category": (data.case_category or "").strip()[:80],
        "tax_item": (data.tax_item or "").strip()[:40],
    }
    if DATABASE_URL and psycopg is not None:
        ensure_feedback_table()
        with psycopg.connect(DATABASE_URL) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO feedback (
                        id, created_at, rating, comment, stage, case_category, tax_item
                    ) VALUES (
                        %(id)s, NOW(), %(rating)s, %(comment)s, %(stage)s, %(case_category)s, %(tax_item)s
                    )
                    """,
                    record,
                )
    else:
        with open(FEEDBACK_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    return record


def ensure_feedback_table():
    if not DATABASE_URL or psycopg is None:
        return
    with psycopg.connect(DATABASE_URL) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS feedback (
                    id uuid PRIMARY KEY,
                    created_at timestamptz NOT NULL DEFAULT NOW(),
                    rating integer NOT NULL CHECK (rating BETWEEN 1 AND 5),
                    comment text NOT NULL DEFAULT '',
                    stage text NOT NULL DEFAULT '',
                    case_category text NOT NULL DEFAULT '',
                    tax_item text NOT NULL DEFAULT ''
                )
                """
            )


def feedback_stats():
    empty = {
        "total": 0,
        "average_rating": None,
        "rating_counts": {"1": 0, "2": 0, "3": 0, "4": 0, "5": 0},
        "by_stage": [],
        "by_tax_item": [],
        "storage": "postgres" if DATABASE_URL and psycopg is not None else "jsonl",
    }

    if DATABASE_URL and psycopg is not None:
        ensure_feedback_table()
        with psycopg.connect(DATABASE_URL, row_factory=dict_row) as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT COUNT(*) total, AVG(rating)::float average_rating FROM feedback")
                summary = cur.fetchone()
                cur.execute("SELECT rating, COUNT(*) count FROM feedback GROUP BY rating ORDER BY rating")
                counts = {str(i): 0 for i in range(1, 6)}
                for row in cur.fetchall():
                    counts[str(row["rating"])] = row["count"]
                cur.execute("SELECT stage, COUNT(*) count, AVG(rating)::float average_rating FROM feedback GROUP BY stage ORDER BY count DESC")
                by_stage = cur.fetchall()
                cur.execute("SELECT tax_item, COUNT(*) count, AVG(rating)::float average_rating FROM feedback GROUP BY tax_item ORDER BY count DESC")
                by_tax_item = cur.fetchall()
        return {
            "total": summary["total"],
            "average_rating": summary["average_rating"],
            "rating_counts": counts,
            "by_stage": by_stage,
            "by_tax_item": by_tax_item,
            "storage": "postgres",
        }

    if not os.path.exists(FEEDBACK_PATH):
        return empty

    records = []
    with open(FEEDBACK_PATH, "r", encoding="utf-8") as f:
        for line in f:
            try:
                records.append(json.loads(line))
            except Exception:
                pass
    if not records:
        return empty

    counts = {str(i): 0 for i in range(1, 6)}
    by_stage_map = {}
    by_tax_map = {}
    total_rating = 0
    for record in records:
        rating = int(record.get("rating", 0))
        if 1 <= rating <= 5:
            counts[str(rating)] += 1
            total_rating += rating
        stage = record.get("stage", "")
        tax_item = record.get("tax_item", "")
        by_stage_map.setdefault(stage, []).append(rating)
        by_tax_map.setdefault(tax_item, []).append(rating)

    def summarize(group):
        rows = []
        for key, ratings in group.items():
            valid = [r for r in ratings if 1 <= r <= 5]
            rows.append({
                "name": key,
                "count": len(valid),
                "average_rating": sum(valid) / len(valid) if valid else None,
            })
        return sorted(rows, key=lambda row: row["count"], reverse=True)

    return {
        "total": len(records),
        "average_rating": total_rating / len(records),
        "rating_counts": counts,
        "by_stage": summarize(by_stage_map),
        "by_tax_item": summarize(by_tax_map),
        "storage": "jsonl",
    }


CASE_CATEGORY_CODES = {
    "1": "稅捐爭議溝通與協調",
    "2": "申訴或陳情",
    "3": "行政救濟諮詢與協助",
}

TAX_ITEM_CODES = {
    "1": "地價稅",
    "2": "使用牌照稅",
    "3": "房屋稅",
    "4": "娛樂稅",
    "5": "印花稅",
    "6": "土地增值稅",
    "7": "契稅",
    "8": "其他類型",
    "land": "地價稅",
    "land_value_increment": "土地增值稅",
    "vehicle": "使用牌照稅",
    "house": "房屋稅",
    "deed": "契稅",
    "entertainment": "娛樂稅",
    "stamp": "印花稅",
    "other": "其他類型",
}

APPLY_METHOD_CODES = {
    "1": "現場申請",
    "2": "書面或傳真申請",
}

REPLY_METHOD_CODES = {
    "1": "現場答復",
    "2": "公文",
    "3": "電話",
}

NOTIFY_METHOD_CODES = {
    "1": "電子郵件",
    "2": "簡訊",
    "3": "無須通知",
}


@app.post("/generate-word")
def generate_word(data: GenerateRequest):

    try:

        doc = Document(TEMPLATE_PATH)

        replacements = {

            "{{apply_year}}": data.apply_year,
            "{{apply_month}}": data.apply_month,
            "{{apply_day}}": data.apply_day,

            "{{formal_statement}}": data.formal_statement,

            "{{case_category_suggestion}}": data.case_category_suggestion,
            "{{tax_items_suggestion}}": data.tax_items_suggestion,

            "{{apply_method_suggestion}}": data.apply_method_suggestion,
            "{{reply_method_suggestion}}": data.reply_method_suggestion,
            "{{notify_method_suggestion}}": data.notify_method_suggestion,

            "{{notify_email}}": data.notify_email,

            "{{representative_name}}": data.representative_name,
            "{{representative_address}}": data.representative_address,
            "{{representative_phone}}": data.representative_phone,

            "{{agent_name}}": data.agent_name,
            "{{agent_address}}": data.agent_address,
            "{{agent_phone}}": data.agent_phone,

            "{{evidence_list}}": data.evidence_list
        }

        replace_all(doc, replacements)

        if data.notify_email.strip() == "":
            remove_email_label(doc)

        if has_placeholder(doc):

            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "message": "產製失敗",
                    "reason": "Word 殘留 placeholder"
                }
            )

        filename = f"申請書_{data.apply_year}{int(data.apply_month):02d}{int(data.apply_day):02d}.docx"

        unique_name = f"{uuid.uuid4()}_{filename}"

        file_path = os.path.join(OUTPUT_DIR, unique_name)

        doc.save(file_path)

        token = str(uuid.uuid4())

        download_store[token] = {
            "path": file_path,
            "filename": filename,
            "expire_time": time.time() + 600,
            "used": False
        }

        return JSONResponse(
            content={
                "success": True,
                "filename": filename,
                "download_url": f"https://tax-word-system-production.up.railway.app/download/{token}"
            }
        )

    except Exception as e:

        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "message": "產製失敗",
                "reason": str(e)
            }
        )


@app.post("/feedback")
def submit_feedback(data: FeedbackRequest):

    try:
        record = write_feedback(data)
        return JSONResponse(content={"success": True, "id": record["id"]})
    except Exception as e:
        return JSONResponse(
            status_code=400,
            content={"success": False, "message": "feedback rejected", "reason": str(e)}
        )


@app.get("/feedback-stats")
def get_feedback_stats():
    return JSONResponse(content=feedback_stats())


@app.get("/feedback-admin", response_class=HTMLResponse)
def feedback_admin_page() -> HTMLResponse:
    stats = feedback_stats()

    def esc(value) -> str:
        return str(value if value is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    total = int(stats.get("total") or 0)
    average = stats.get("average_rating")
    average_text = f"{float(average):.2f}" if average is not None else "-"
    storage = stats.get("storage") or "-"
    rating_counts = stats.get("rating_counts") or {}
    max_count = max([int(rating_counts.get(str(i), 0) or 0) for i in range(1, 6)] + [1])

    rating_rows = []
    for score in range(5, 0, -1):
        count = int(rating_counts.get(str(score), 0) or 0)
        width = round(count / max_count * 100) if max_count else 0
        rating_rows.append(
            f"""
            <div class="bar-row">
              <div class="bar-label">{score} 分</div>
              <div class="bar-track"><div class="bar-fill" style="width:{width}%"></div></div>
              <div class="bar-count">{count}</div>
            </div>
            """
        )

    def table_rows(rows, empty_label):
        if not rows:
            return f'<tr><td colspan="3" class="empty">{empty_label}</td></tr>'
        rendered = []
        for row in rows:
            name = row.get("stage") or row.get("tax_item") or row.get("name") or "未標示"
            count = int(row.get("count") or 0)
            avg = row.get("average_rating")
            avg_text = f"{float(avg):.2f}" if avg is not None else "-"
            rendered.append(
                f"<tr><td>{esc(name)}</td><td>{count}</td><td>{avg_text}</td></tr>"
            )
        return "\n".join(rendered)

    html = f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>滿意度統計 / Feedback Dashboard</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Noto Sans TC',Arial,sans-serif;background:#eef3f7;color:#243241;padding:24px;min-height:100vh}}
.wrap{{max-width:980px;margin:0 auto}}
.top{{display:flex;align-items:flex-end;justify-content:space-between;gap:1rem;margin-bottom:1rem}}
h1{{font-size:1.55rem;line-height:1.35;color:#1f2d3d}}
.subtitle{{color:#607080;font-size:.92rem;margin-top:.3rem}}
.updated{{color:#607080;font-size:.85rem;text-align:right}}
.cards{{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:12px;margin-bottom:12px}}
.card,.panel{{background:#fff;border:1px solid #dce5ee;border-radius:8px;box-shadow:0 1px 4px rgba(20,40,60,.05)}}
.card{{padding:1rem}}
.card-title{{font-size:.86rem;color:#607080;margin-bottom:.45rem}}
.card-value{{font-size:1.9rem;font-weight:800;color:#1f2d3d}}
.card-note{{font-size:.8rem;color:#7b8895;margin-top:.35rem}}
.grid{{display:grid;grid-template-columns:minmax(280px,1fr) minmax(280px,1fr);gap:12px}}
.panel{{padding:1rem;margin-bottom:12px}}
h2{{font-size:1rem;color:#1f2d3d;margin-bottom:.8rem}}
.bar-row{{display:grid;grid-template-columns:48px 1fr 48px;align-items:center;gap:.7rem;margin:.65rem 0}}
.bar-label,.bar-count{{font-size:.9rem;color:#526170}}
.bar-count{{text-align:right;font-weight:700;color:#1f2d3d}}
.bar-track{{height:12px;background:#e8eef4;border-radius:999px;overflow:hidden}}
.bar-fill{{height:100%;background:#2f80b9;border-radius:999px;min-width:0}}
table{{width:100%;border-collapse:collapse;font-size:.9rem}}
th,td{{padding:.65rem .55rem;border-bottom:1px solid #edf1f5;text-align:left;vertical-align:top}}
th{{color:#526170;font-weight:700;background:#f8fafc}}
td:nth-child(2),td:nth-child(3),th:nth-child(2),th:nth-child(3){{text-align:right;width:86px}}
.empty{{color:#8492a0;text-align:center!important;padding:1.2rem}}
.foot{{font-size:.82rem;color:#6d7b88;line-height:1.6;margin-top:.4rem}}
@media(max-width:760px){{
  body{{padding:16px}}
  .top{{display:block}}
  .updated{{text-align:left;margin-top:.5rem}}
  .cards,.grid{{grid-template-columns:1fr}}
}}
</style>
</head>
<body>
<main class="wrap">
  <div class="top">
    <div>
      <h1>滿意度統計<br>Feedback Dashboard</h1>
      <div class="subtitle">納保申請助理服務回饋統計</div>
    </div>
    <div class="updated">資料來源：{esc(storage)}<br>重新整理頁面即可更新</div>
  </div>

  <section class="cards">
    <div class="card">
      <div class="card-title">總填寫筆數 / Total</div>
      <div class="card-value">{total}</div>
    </div>
    <div class="card">
      <div class="card-title">平均分數 / Average</div>
      <div class="card-value">{average_text}</div>
      <div class="card-note">滿分 5 分</div>
    </div>
    <div class="card">
      <div class="card-title">儲存狀態 / Storage</div>
      <div class="card-value" style="font-size:1.35rem">{esc(storage)}</div>
      <div class="card-note">postgres 代表長期資料庫</div>
    </div>
  </section>

  <section class="panel">
    <h2>分數分布 / Rating Counts</h2>
    {''.join(rating_rows)}
  </section>

  <section class="grid">
    <div class="panel">
      <h2>依階段統計 / By Stage</h2>
      <table>
        <thead><tr><th>階段</th><th>筆數</th><th>平均</th></tr></thead>
        <tbody>{table_rows(stats.get("by_stage") or [], "尚無階段資料")}</tbody>
      </table>
    </div>
    <div class="panel">
      <h2>依稅目統計 / By Tax Item</h2>
      <table>
        <thead><tr><th>稅目</th><th>筆數</th><th>平均</th></tr></thead>
        <tbody>{table_rows(stats.get("by_tax_item") or [], "尚無稅目資料")}</tbody>
      </table>
    </div>
  </section>
  <p class="foot">此頁只顯示彙總統計，不顯示個別建議內容。若使用者誤填個資，不會直接出現在此頁面。</p>
</main>
</body>
</html>"""
    return HTMLResponse(content=html)


@app.get("/survey", response_class=HTMLResponse)
def survey_page(
    stage: str = Query(default="statement"),
    c: str = Query(default=""),
    t: str = Query(default=""),
    case_category: str = Query(default=""),
    tax_item: str = Query(default=""),
) -> HTMLResponse:

    case_category = case_category or normalize_choice(c, CASE_CATEGORY_CODES)
    tax_item = tax_item or normalize_choice(t, TAX_ITEM_CODES)

    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")

    html = f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>納保申請助理滿意度調查 / Satisfaction Survey</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Noto Sans TC',sans-serif;background:#f0f4f8;padding:20px;min-height:100vh;display:flex;align-items:center;justify-content:center}}
.box{{width:100%;max-width:560px;background:#fff;border-radius:12px;padding:2rem;box-shadow:0 2px 12px rgba(0,0,0,.1)}}
h1{{font-size:1.25rem;color:#2c3e50;margin-bottom:.6rem}}
h1 span{{font-size:1rem;color:#5d6d7e}}
p{{color:#666;font-size:.9rem;line-height:1.7;margin-bottom:1rem}}
.ratings{{display:grid;grid-template-columns:repeat(5,1fr);gap:.5rem;margin:1rem 0}}
.ratings button{{border:1px solid #b8c7d9;background:#fff;color:#2c3e50;border-radius:8px;padding:.8rem .2rem;font-size:1rem;cursor:pointer}}
.ratings button.selected{{background:#2980b9;color:#fff;border-color:#2980b9}}
textarea{{width:100%;height:120px;resize:vertical;border:1px solid #ccc;border-radius:8px;padding:.75rem;font-family:inherit;font-size:.95rem;line-height:1.5}}
.submit{{width:100%;margin-top:1rem;border:0;border-radius:8px;background:#2980b9;color:#fff;padding:.85rem;font-size:1rem;font-weight:600;cursor:pointer}}
.submit:disabled{{background:#aaa;cursor:not-allowed}}
#msg{{margin-top:.8rem;font-size:.9rem}}
.ok{{color:#1e8449}}.err{{color:#c0392b}}
</style>
</head>
<body>
<div class="box">
  <h1>納保申請助理滿意度調查<br><span>Satisfaction Survey</span></h1>
  <p>請針對本次服務給予 1 至 5 分評分，也可以留下建議事項。請勿填寫姓名、電話、地址、車牌或其他個人資料。</p>
  <p>Please rate this service from 1 to 5 and optionally leave suggestions. Do not enter your name, phone number, address, license plate number, or other personal information.</p>
  <div class="ratings" id="ratings">
    <button type="button" data-rating="1">1</button>
    <button type="button" data-rating="2">2</button>
    <button type="button" data-rating="3">3</button>
    <button type="button" data-rating="4">4</button>
    <button type="button" data-rating="5">5</button>
  </div>
  <textarea id="comment" placeholder="建議事項（選填） / Suggestions (optional)"></textarea>
  <button class="submit" id="submit" type="button" disabled>送出 / Submit</button>
  <div id="msg"></div>
</div>
<script>
let rating = 0;
const stage = {json.dumps(stage, ensure_ascii=False)};
const caseCategory = {json.dumps(case_category, ensure_ascii=False)};
const taxItem = {json.dumps(tax_item, ensure_ascii=False)};
document.querySelectorAll('[data-rating]').forEach(btn => {{
  btn.addEventListener('click', () => {{
    rating = Number(btn.dataset.rating);
    document.querySelectorAll('[data-rating]').forEach(b => b.classList.toggle('selected', b === btn));
    document.getElementById('submit').disabled = false;
  }});
}});
document.getElementById('submit').addEventListener('click', async () => {{
  const submit = document.getElementById('submit');
  const msg = document.getElementById('msg');
  submit.disabled = true;
      msg.textContent = '送出中... / Submitting...';
  msg.className = '';
  try {{
    const res = await fetch('/feedback', {{
      method: 'POST',
      headers: {{'Content-Type': 'application/json'}},
      body: JSON.stringify({{
        rating,
        comment: document.getElementById('comment').value,
        stage,
        case_category: caseCategory,
        tax_item: taxItem
      }})
    }});
    const json = await res.json();
    if (json.success) {{
      msg.textContent = '感謝您的回饋。 / Thank you for your feedback.';
      msg.className = 'ok';
    }} else {{
      throw new Error(json.reason || '送出失敗');
    }}
  }} catch (e) {{
    msg.textContent = '送出失敗，請稍後再試。 / Submission failed. Please try again later.';
    msg.className = 'err';
    submit.disabled = false;
  }}
}});
</script>
</body>
</html>"""
    return HTMLResponse(content=html)


@app.get("/form", response_class=HTMLResponse)
def form_page(
    apply_year: str = Query(default=""),
    apply_month: str = Query(default=""),
    apply_day: str = Query(default=""),
    case_category_suggestion: str = Query(default=""),
    tax_items_suggestion: str = Query(default=""),
    apply_method_suggestion: str = Query(default=""),
    reply_method_suggestion: str = Query(default=""),
    notify_method_suggestion: str = Query(default=""),
    formal_statement: str = Query(default=""),
    y: str = Query(default=""),
    m: str = Query(default=""),
    d: str = Query(default=""),
    c: str = Query(default=""),
    t: str = Query(default=""),
    a: str = Query(default=""),
    r: str = Query(default=""),
    n: str = Query(default=""),
    s: str = Query(default=""),
) -> HTMLResponse:

    apply_year = apply_year or y
    apply_month = apply_month or m
    apply_day = apply_day or d
    case_category_suggestion = case_category_suggestion or normalize_choice(c, CASE_CATEGORY_CODES)
    tax_items_suggestion = tax_items_suggestion or normalize_choice(t, TAX_ITEM_CODES)
    apply_method_suggestion = apply_method_suggestion or normalize_choice(a, APPLY_METHOD_CODES)
    reply_method_suggestion = reply_method_suggestion or normalize_choice(r, REPLY_METHOD_CODES)
    notify_method_suggestion = notify_method_suggestion or normalize_choice(n, NOTIFY_METHOD_CODES)
    formal_statement = formal_statement or s

    # 單次使用邏輯：只有在帶有實質參數時啟用
    has_params = bool(apply_year or formal_statement or case_category_suggestion)

    if has_params:
        param_str = "|".join([
            apply_year, apply_month, apply_day,
            case_category_suggestion, tax_items_suggestion,
            apply_method_suggestion, reply_method_suggestion,
            notify_method_suggestion, formal_statement,
        ])
        url_hash = hashlib.sha256(param_str.encode("utf-8")).hexdigest()

        now = time.time()

        if url_hash in form_access_store:
            if now <= form_access_store[url_hash]:
                # 30 分鐘內第二次點擊 → 失效
                return HTMLResponse(content=EXPIRED_HTML, status_code=410)
            else:
                # 超過 30 分鐘 → 允許重新使用
                del form_access_store[url_hash]

        # 首次存取：登記，有效封鎖期 30 分鐘
        form_access_store[url_hash] = now + 1800

    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")

    html = f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>納保申請書產製</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Noto Sans TC',sans-serif;background:#f0f4f8;padding:20px;min-height:100vh}}
.container{{max-width:680px;margin:0 auto;background:#fff;border-radius:12px;padding:2rem;box-shadow:0 2px 12px rgba(0,0,0,.1)}}
h1{{font-size:1.3rem;color:#2c3e50;border-bottom:2px solid #3498db;padding-bottom:.6rem;margin-bottom:1.2rem}}
.field{{margin-bottom:1rem}}
label{{display:block;font-weight:600;margin-bottom:.3rem;color:#444;font-size:.9rem}}
input,textarea{{width:100%;padding:.55rem .75rem;border:1px solid #ccc;border-radius:6px;font-size:.95rem;font-family:inherit}}
input:focus,textarea:focus{{outline:none;border-color:#3498db;box-shadow:0 0 0 2px rgba(52,152,219,.2)}}
textarea{{height:200px;resize:vertical;line-height:1.6}}
.row{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:.8rem}}
.hint{{font-size:.8rem;color:#888;margin-top:.25rem}}
.hint.warn{{color:#e67e22;font-weight:600}}
button{{background:#2980b9;color:#fff;border:none;padding:.85rem;border-radius:8px;font-size:1rem;cursor:pointer;width:100%;margin-top:.5rem;font-family:inherit;font-weight:600;letter-spacing:.03em}}
button:hover{{background:#2471a3}}
button:disabled{{background:#aaa;cursor:not-allowed}}
#spinner{{display:none;text-align:center;padding:.8rem;color:#555}}
.error{{color:#c0392b;background:#fdf0f0;border:1px solid #f5c6cb;padding:.75rem;border-radius:6px;margin-top:.8rem}}
.success{{color:#1e8449;background:#eafaf1;border:1px solid #a9dfbf;padding:.75rem;border-radius:6px;margin-top:.8rem}}
.divider{{border:none;border-top:1px solid #eee;margin:1.2rem 0}}
.feedback{{display:none;margin-top:1rem;border:1px solid #d7e3ef;border-radius:10px;padding:1rem;background:#f8fbff}}
.feedback-title{{font-weight:700;color:#2c3e50;margin-bottom:.35rem}}
.feedback-note{{font-size:.82rem;color:#777;margin-bottom:.75rem;line-height:1.5}}
.rating-row{{display:grid;grid-template-columns:repeat(5,1fr);gap:.4rem;margin-bottom:.75rem}}
.rating-row button{{margin:0;background:#fff;color:#2c3e50;border:1px solid #b8c7d9;padding:.6rem .2rem}}
.rating-row button.selected{{background:#2980b9;color:#fff;border-color:#2980b9}}
#feedback_comment{{height:90px}}
#feedback_msg{{font-size:.85rem;margin-top:.5rem}}
</style>
</head>
<body>
<div class="container">
  <h1>📄 納保申請書產製</h1>
  <p style="color:#666;font-size:.88rem;margin-bottom:1.2rem">由納保申請助理轉介。所有欄位已自動填入，確認後點「產製申請書」即可下載。</p>

  <div class="row">
    <div class="field">
      <label>申請年份（民國）</label>
      <input id="apply_year" value="{esc(apply_year)}" placeholder="114">
    </div>
    <div class="field">
      <label>月份</label>
      <input id="apply_month" value="{esc(apply_month)}" placeholder="3">
    </div>
    <div class="field">
      <label>日期</label>
      <input id="apply_day" value="{esc(apply_day)}" placeholder="15">
    </div>
  </div>

  <div class="field">
    <label>案件類型建議</label>
    <input id="case_category_suggestion" value="{esc(case_category_suggestion)}">
  </div>

  <div class="field">
    <label>稅目建議</label>
    <input id="tax_items_suggestion" value="{esc(tax_items_suggestion)}">
  </div>

  <hr class="divider">

  <div class="field">
    <label>申請事由</label>
    <textarea id="formal_statement" placeholder="申請事由">{esc(formal_statement)}</textarea>
    <div class="hint">※ 如需修改可直接編輯</div>
  </div>

  <hr class="divider">

  <div class="field">
    <label>申請方式</label>
    <input id="apply_method_suggestion" value="{esc(apply_method_suggestion)}">
  </div>
  <div class="field">
    <label>回覆方式</label>
    <input id="reply_method_suggestion" value="{esc(reply_method_suggestion)}">
  </div>
  <div class="field">
    <label>通知方式</label>
    <input id="notify_method_suggestion" value="{esc(notify_method_suggestion)}">
  </div>

  <button id="btn" onclick="generate()">產製申請書</button>
  <div id="spinner">⏳ 正在產製，請稍候…</div>
  <div id="result"></div>
  <div class="feedback" id="feedback">
    <div class="feedback-title">滿意度調查 / Satisfaction Survey</div>
    <div class="feedback-note">請針對本次申請書產製服務給予 1 至 5 分評分，也可以留下建議事項。請勿填寫姓名、電話、地址、車牌或其他個人資料。<br>Please rate this form generation service from 1 to 5 and optionally leave suggestions. Do not enter your name, phone number, address, license plate number, or other personal information.</div>
    <div class="rating-row" id="feedback_ratings">
      <button type="button" data-feedback-rating="1">1</button>
      <button type="button" data-feedback-rating="2">2</button>
      <button type="button" data-feedback-rating="3">3</button>
      <button type="button" data-feedback-rating="4">4</button>
      <button type="button" data-feedback-rating="5">5</button>
    </div>
    <textarea id="feedback_comment" placeholder="建議事項（選填） / Suggestions (optional)"></textarea>
    <button id="feedback_submit" type="button" disabled onclick="submitFeedback()">送出回饋 / Submit Feedback</button>
    <div id="feedback_msg"></div>
  </div>
</div>

<script>
let feedbackRating = 0;
document.querySelectorAll('[data-feedback-rating]').forEach(btn => {{
  btn.addEventListener('click', () => {{
    feedbackRating = Number(btn.dataset.feedbackRating);
    document.querySelectorAll('[data-feedback-rating]').forEach(b => b.classList.toggle('selected', b === btn));
    document.getElementById('feedback_submit').disabled = false;
  }});
}});

async function generate() {{
  const btn = document.getElementById('btn');
  const spinner = document.getElementById('spinner');
  const result = document.getElementById('result');

  const formal = document.getElementById('formal_statement').value.trim();
  if (!formal) {{
    result.innerHTML = '<div class="error">⚠️ 申請事由為空，請確認後再試！</div>';
    return;
  }}

  btn.disabled = true;
  spinner.style.display = 'block';
  result.innerHTML = '';

  const payload = {{
    apply_year: document.getElementById('apply_year').value.trim(),
    apply_month: document.getElementById('apply_month').value.trim(),
    apply_day: document.getElementById('apply_day').value.trim(),
    formal_statement: formal,
    case_category_suggestion: document.getElementById('case_category_suggestion').value.trim(),
    tax_items_suggestion: document.getElementById('tax_items_suggestion').value.trim(),
    apply_method_suggestion: document.getElementById('apply_method_suggestion').value.trim(),
    reply_method_suggestion: document.getElementById('reply_method_suggestion').value.trim(),
    notify_method_suggestion: document.getElementById('notify_method_suggestion').value.trim(),
    notify_email: '',
    evidence_list: ''
  }};

  try {{
    const res = await fetch('/generate-word', {{
      method: 'POST',
      headers: {{'Content-Type': 'application/json'}},
      body: JSON.stringify(payload)
    }});
    const json = await res.json();
    if (json.success && json.download_url) {{
      result.innerHTML = '<div class="success">✅ 申請書已產製完成！若下載未自動開始，<a href="' + json.download_url + '" target="_blank">請點此下載</a></div>';
      document.getElementById('feedback').style.display = 'block';
      setTimeout(() => {{ window.location.href = json.download_url; }}, 500);
    }} else {{
      result.innerHTML = '<div class="error">❌ 產製失敗：' + (json.reason || json.message || '未知錯誤') + '</div>';
    }}
  }} catch(e) {{
    result.innerHTML = '<div class="error">❌ 網路錯誤：' + e.message + '</div>';
  }} finally {{
    btn.disabled = false;
    spinner.style.display = 'none';
  }}
}}

async function submitFeedback() {{
  const submit = document.getElementById('feedback_submit');
  const msg = document.getElementById('feedback_msg');
  submit.disabled = true;
  msg.textContent = '送出中... / Submitting...';
  msg.className = '';
  try {{
    const res = await fetch('/feedback', {{
      method: 'POST',
      headers: {{'Content-Type': 'application/json'}},
      body: JSON.stringify({{
        rating: feedbackRating,
        comment: document.getElementById('feedback_comment').value,
        stage: 'form_generated',
        case_category: document.getElementById('case_category_suggestion').value.trim(),
        tax_item: document.getElementById('tax_items_suggestion').value.trim()
      }})
    }});
    const json = await res.json();
    if (json.success) {{
      msg.textContent = '感謝您的回饋。 / Thank you for your feedback.';
      msg.style.color = '#1e8449';
    }} else {{
      throw new Error(json.reason || '送出失敗');
    }}
  }} catch(e) {{
    msg.textContent = '送出失敗，請稍後再試。 / Submission failed. Please try again later.';
    msg.style.color = '#c0392b';
    submit.disabled = false;
  }}
}}
</script>
</body>
</html>"""
    return HTMLResponse(content=html)


@app.get("/download/{token}")
def download_file(token: str):

    if token not in download_store:
        return JSONResponse(status_code=404, content={"message": "file not found"})

    record = download_store[token]

    if time.time() > record["expire_time"]:

        if os.path.exists(record["path"]):
            os.remove(record["path"])

        del download_store[token]

        return JSONResponse(status_code=410, content={"message": "link expired"})

    if record["used"]:
        return JSONResponse(status_code=410, content={"message": "file already downloaded"})

    record["used"] = True

    return FileResponse(
        record["path"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=record["filename"]
    )


async def cleanup_expired():

    while True:

        now = time.time()

        # 清理 download_store
        expired_downloads = [
            token for token, record in download_store.items()
            if now > record["expire_time"]
        ]
        for token in expired_downloads:
            record = download_store[token]
            if os.path.exists(record["path"]):
                os.remove(record["path"])
            del download_store[token]

        # 清理 form_access_store
        expired_forms = [
            h for h, expire_time in form_access_store.items()
            if now > expire_time
        ]
        for h in expired_forms:
            del form_access_store[h]

        await asyncio.sleep(60)


@app.on_event("startup")
async def startup_event():

    asyncio.create_task(cleanup_expired())
